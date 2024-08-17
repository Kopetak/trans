import os
import re
import pandas as pd
from tqdm import tqdm
import docx
from difflib import SequenceMatcher

##################################################
#Step4: マッチングの検索
####################main code#####################
#読み込みファイル
path = os.getcwd()
folder = path + "/2_make_corpus"

#コーパス読み込み
df_corpus = pd.read_excel(path + "/corpus.xlsx")

df = pd.DataFrame()
    
#日ファイル名
name = "gor2404a"
filename =path + "/3_translation/" + name + ".docx"
doc = docx.Document(filename)

counter = 0

for pr in range(len(doc.paragraphs)):
    parag = doc.paragraphs[pr].text.replace("　", "").replace(" ", "").replace('\n',"")
    if len(parag) > 0:
        df = df.append([parag])
        counter = counter + 1
df.index = range(counter)
df = df.dropna(how = "all")
   
pn = len(df)  # パラグラフ数
   
for p in range(pn):
       
       #カット
       para_JP = re.split(r'(?<=。)', str(df.iloc[p,0]))
       df_JP = pd.DataFrame(para_JP)
       
       #スペースの削除
       df_JP = df_JP.replace("　", "").replace(" ", "")
       df_JP = df_JP[lambda df_JP: df_JP.iloc[:,0].str.len() > 0]
   
       df_JP.columns = ["JP"]
       
       #結合
       if p == 0:
           df_out = df_JP
       else:
           df_out = pd.concat([df_out,df_JP], axis = 0, ignore_index=True)
           
           
#準備
df_out["類似文"] = "NaN" #1
df_out["類似度"] = "NaN" #2
df_out["出所"] = "NaN" #3
df_out["EN"] = "NaN" #4

for k in tqdm(range(len(df_out))):
    src = df_out.iloc[k,0]
    s_len = len(src)
    
    rmax = 0
    #ゲシュタルトパターン・マッチングによる部分一致比較
    for i in range(len(df_corpus)):
        trg = df_corpus.iloc[i,0] 
        t_len = len(trg)
    
        if s_len <= t_len:
            r = max([SequenceMatcher(None, src, trg[i:i+s_len]).ratio() for i in range(t_len-s_len+1)])
        # else:
        #     r = max([SequenceMatcher(None, src[i:i+t_len], trg).ratio() for i in range(s_len-t_len+1)])
        
        if r > rmax:
            rmax = r
            imax = i
                
        if rmax == 1:
            break
                
    df_out.iloc[k,1] = df_corpus.iloc[imax,0]
    df_out.iloc[k,2] = rmax
    df_out.iloc[k,3] = df_corpus.iloc[imax,2]
    df_out.iloc[k,4] = df_corpus.iloc[imax,1]
    
#出力
filename = path + "/" + name + "_result.xlsx"
df_out.to_excel(filename,index = False, encoding = "shift-jis")