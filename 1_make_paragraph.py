import os
import pandas as pd
from tqdm import tqdm
import docx

##################################################
#Step1: 与えたwordファイル（日英）から段落ごとの対応表を作成
# →　エクセルに出力したあと、人力で整合的になるよう並び替え
####################main code#####################
#読み込みファイル
path = os.getcwd()
files = os.listdir(path + "/0_input")
files = [s for s in files if 'JP' in s]

#データフレームの準備
master = pd.DataFrame()

for file in tqdm(files):
    
    testdf = pd.DataFrame()
    testdf_EN = pd.DataFrame()
    
    #日ファイル名
    filename =path + "/0_input/" + file
    doc = docx.Document(filename)
    
    #英ファイル名
    filename_EN = filename.replace("JP", "EN")
    doc_EN = docx.Document(filename_EN)
    
    counter = 0
    
    for pr in range(len(doc.paragraphs)):
        parag = doc.paragraphs[pr].text.replace("　", "").replace(" ", "").replace('\n',"")
        if len(parag) > 0:
            testdf = testdf.append([parag])
            counter = counter + 1
    testdf.index = range(counter)

    counter = 0
        
    for pr in range(len(doc_EN.paragraphs)):
        parag = doc_EN.paragraphs[pr].text.replace("　", "").replace('\n',"")
        if len(parag) > 1:
            testdf_EN = testdf_EN.append([parag])
            counter = counter + 1
    testdf_EN.index = range(counter)
        
    master = pd.concat([testdf, testdf_EN], axis=1, ignore_index=True)
    master.columns = ["JP", "EN"]
    outname = file.replace("_JP", "").replace(".docx", "")
    master["file"] = outname

master.to_excel(path + "/1_make_paragraph/raw/" + outname + "_para.xlsx", index = False, encoding = "shift-jis")